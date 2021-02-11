from docx.enum.text import WD_BREAK

from documents.builder.styling import (add_styling,
                                       format_table)

from documents.builder.content import (add_content, get_text,
                                       get_display_list)
from documents.builder.utils import (is_fee_item, is_fee_table,
                                     get_num_cols)
from documents.builder.structure import (full_width_content, skip_row,
                                         build_agg_total)


def build_fee_tables(all_platforms, structure, table_type, document, style,
                 table_style, content, platform_slice, num_table_splits, orientation):



    for table_split in range(num_table_splits):

        num_cols, platform_slice = split_tables(num_table_splits, table_split,
                                           all_platforms, structure, table_type,
                                           orientation, platform_slice)

        table = document.add_table(rows=0, cols=num_cols)
        format_table(table, table_style, style)
        build_rows(table, table_type, structure, platform_slice, content, all_platforms)

        if table_split != num_table_splits-1:
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)


def split_tables(num_table_splits, table_split, all_platforms, structure,
                 table_type, orientation, platform_slice):

    if "LANDSCAPE" in str(orientation):
        overflow_limit = structure.landscape_overflow_limit
    elif "PORTRAIT" in str(orientation):
        overflow_limit = structure.portrait_overflow_limit

    if num_table_splits <= 1:
        num_cols = get_num_cols(platform_slice, structure, table_type,
                                num_table_splits)
    else:
        if table_split == num_table_splits-1:
            # Partial table
            num_cols = (len(all_platforms) % overflow_limit) + 1
            start = (table_split) * overflow_limit
            end = start + num_cols
            platform_slice = all_platforms[start:end]
        else:
            # Full table
            num_cols = overflow_limit + 1
            start = (table_split) * overflow_limit
            end = start + overflow_limit
            platform_slice = all_platforms[start:end]

    return num_cols, platform_slice


def build_rows(table, table_type, structure, platform_slice, content_list, all_platforms):
    '''
    Calls functions to create rows/cells.
    '''
    try:
        # Checks the given table is a fee table
        if is_fee_table(table_type):
            for content in content_list:
                if content.type == "Aggregated Total":
                    build_agg_total(table, platform_slice, content)
                elif not is_fee_item(content):
                    print(1)
                    # Builds cells for non-fee items
                    build_cells(table, platform_slice, content, content.type, structure, all_platforms)
                    print(13)
                else:
                    # Builds cells for fee items
                    display_list = get_display_list(content, platform_slice)
                    for property, display_name in display_list:
                        build_cells(table, platform_slice, content, property, structure, all_platforms, display=display_name)
    except Exception as e:
        print(e)
        print("HG")

def build_cells(table, platform_slice, content, description, structure, all_platforms, **kwargs):
    '''
    Creates rows and cells and calls functions to add text and styling
    '''
    text_list, data_list = get_text(platform_slice, description, is_fee_item(content), content, structure, all_platforms)
    if not skip_row(data_list, structure):
        row_cells = table.add_row().cells
        display_name = kwargs.get("display")

        if not display_name:
            display_name = content

        element = content.element
        try:
            for (i, cell) in enumerate(row_cells):

                if not full_width_content(content, structure):
                    add_content(cell, i, display_name, text_list, element)
                    add_styling(cell, i, element)

                else:
                    if i == 0:
                        add_content(cell, i, display_name, text_list, element)
                        add_styling(cell, i, element)
                    else:
                        row_cells[0].merge(cell)

        except Exception as e:
            print(e)
            print("H")
