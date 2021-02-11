from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.shortcuts import get_list_or_404

from documents.models import Element



def add_bgcolor(cell, element):
    '''
    Adds a background (fill) colour to the given cell based on given element.
    '''
    if element.background_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcShd = OxmlElement('w:shd')
        tcShd.set(qn('w:fill'), element.background_color)
        tcPr.append(tcShd)


def is_cell(item):
    return hasattr(item, '_tc')


def append_border_element(item):
    '''
    Appends either a tcBorders ot tblBorders element to the given item (cell or
    table) and returns that element for further modification.
    '''
    # For cells
    if is_cell(item):
        tc = item._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)

    # For tables
    else:
        tblPr = item._tblPr
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)

    return borders


def add_border(item, element):
    '''
    Adds a border to the given item (cell or table) based on a given Element
    (cell) or Style (table).
    '''
    if element.border_sides:

        borders = append_border_element(item)

        sides = element.border_sides.split(",")

        for side in sides:
            tag = f"w:{side}"
            node = borders.find(qn(tag))

            if node is None:
                node = OxmlElement(tag)
                borders.append(node)

            node.set(qn('w:sz'), str(element.border_width*8))
            node.set(qn('w:val'), 'single')
            node.set(qn('w:color'), element.border_color)


def apply_rowheader_alignment(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_vertical_alignment(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styling(cell, i, element):

    assign_cell_style(cell, i, element)
    # print("Cell style added")
    add_bgcolor(cell, element)
    # print("BGcolor added")
    add_border(cell, element)
    add_vertical_alignment(cell)


def assign_cell_style(cell, i, element):
    '''
    Assigns style to the given cell.
    '''
    for paragraph in cell.paragraphs:
        paragraph.style = element

        if i == 0:
            apply_rowheader_alignment(paragraph)

def get_platforms(platforms, structure):
    '''
    Returns the platforms to be used, depending in whether alts are to be
    included.
    '''
    if structure.display_alternate_platform:
        return platforms
    else:
        return [x for x in platforms if x.status != 'Alternative']


def set_color(item, value):
    '''
    Converts the stored hex color to RGB color before assigning to the given
    item's color property.
    '''
    if value:
        rgb = tuple(int(value[i:i+2], 16) for i in (0, 2, 4))
        (r,g,b)= rgb
        value = RGBColor(r,g,b)
        item.color.rgb = value


def get_option_value(element, option, style):
    '''
    This is not a great way to build base styles. If performance is affected,
    consider using base_style property as per python-docx docs.
    '''
    if element:
        value = getattr(element, "font_"+option)

        if value:
            return value
        else:
            return getattr(style, "font_"+option)

    else:
        return getattr(style, "font_"+option)


def apply_font_styling(docstyle, element, style):
    '''
    Applies font styling options to the given element.
    Matches the given style option with a model field of similar name.
    '''
    style_options = ["bold", "color", "italic", "name", "size", "small_caps",
                     "underline"]
    font = docstyle.font
    for option in style_options:
        value = get_option_value(element, option, style)
        if option == 'size':
            value = Pt(value)

        if option == 'color':
            set_color(font, value)
        else:
            setattr(font, option, value)


def add_table_cell_margins(table, **kwargs):
    """
    table:  actual table instance you want to modify

    usage:

        set_table_cell_margins(table, top=50, bottom=50, start=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    """
    tblPr = table._tblPr
    tblCellMar = OxmlElement('w:tblCellMar')

    for m in [
        "top",
        "bottom",
        "start",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            # w:top
            node.set(qn('w:w'), str(kwargs.get(m)))
            # w:w set kwargs[0]
            node.set(qn('w:type'), 'dxa')
            # w:type set 'dxa'
            tblCellMar.append(node)

    tblPr.append(tblCellMar)

def build_element_style_set(document, style):
    '''
    Builds the style set for the elements in the doc.
    Iterates through each element and creates a new style for each one.
    '''

    elements = get_list_or_404(Element, style=style)

    for element in elements:
        document.styles.add_style(element.type, WD_STYLE_TYPE.PARAGRAPH)
        docstyle = document.styles[element]
        apply_font_styling(docstyle, element, style)


def apply_paragraph_styling(table_style):

    if 'TABLE' in str(table_style.type):
        format = table_style.paragraph_format
        format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format.line_spacing = 1.0
        format.space_after = 0


def build_table_style_set(document, style):
    '''
    Builds the style for the tables.
    '''
    table_style = document.styles.add_style('tableStyle', WD_STYLE_TYPE.TABLE)
    element = None
    apply_font_styling(table_style, element, style)
    apply_paragraph_styling(table_style)

    return table_style


def assign_table_style(table, table_style):
    table.style = table_style


def set_row_header_width(table, **kwargs):
    '''
    Sets the first column (row header) width to 4cm
    '''
    table_type = kwargs.get("table_type")

    if table_type == "Product":
        row_header = table.columns[0]
        row_header.width = Inches(1.9)
    else:
        row_header = table.columns[0]
        row_header.width = Inches(1.5748)

def format_table(table, table_style, style, **kwargs):

    table_type = kwargs.get("table_type")


    assign_table_style(table, table_style)
    add_border(table, style)
    add_table_cell_margins(table,
                           top=style.cell_margin_top,
                           bottom=style.cell_margin_bottom,
                           start=style.cell_margin_start,
                           end=style.cell_margin_end)
    if table_type:
        set_row_header_width(table, table_type=table_type)
    else:
        set_row_header_width(table)
