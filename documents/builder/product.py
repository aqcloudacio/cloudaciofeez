from django.shortcuts import get_list_or_404, get_object_or_404
from docx.enum.text import WD_BREAK
from decimal import Decimal

from documents.builder.styling import (build_element_style_set,
                                       build_table_style_set,
                                       add_styling,
                                       format_table)
from documents.builder.content import add_content
from documents.builder.structure import full_width_content
from documents.builder.utils import is_content_item


def build_product_tables(document, style, table_style, structure, content, all_platforms, investments):

    num_cols = get_num_cols(structure)
    table = document.add_table(rows=0, cols=num_cols)

    all_platforms = [p for p in all_platforms if not p.get_destination_platforms()]

    format_table(table, table_style, style, table_type=content[0].table_type)
    build_rows(table, structure, content, all_platforms, investments)

    # p = document.add_paragraph()
    # run = p.add_run()
    # run.add_break(WD_BREAK.PAGE)

def build_rows(table, structure, content_list, all_platforms, investments):
    '''
    Calls functions to create rows/cells.
    '''
    col_headers = get_column_headers(structure)
    all_platforms = remove_alternatives(all_platforms)

    col_header = get_object_or_404(content_list, type="Description")
    grand_total = get_object_or_404(content_list, type="Total")

    try:
        #Build row header
        build_cells(table, col_header, structure, all_platforms, is_content_item(col_header), col_headers)

        for platform in all_platforms:
            for content in [c for c in content_list if (c.type != "Total") and (c.type != "Description")]:
                if not is_content_item(content):
                    build_cells(table, content, structure, all_platforms, is_content_item(content), col_headers, row_item=platform)
                else:
                    #build investment rows
                    inv_list = [inv for inv in investments if inv.platform == platform]
                    inv_list += platform.full_selldowns
                    for inv in inv_list:
                        build_cells(table, content, structure, all_platforms, is_content_item(content), col_headers, row_item=inv)
        #Build grand total
        build_cells(table, grand_total, structure, all_platforms, is_content_item(grand_total), col_headers, row_item=all_platforms)



    except Exception as e:
        print(e)


def build_cells(table, content, structure, all_platforms, is_content_item, col_headers, **kwargs):
    '''
    Creates rows and cells and calls functions to add text and styling
    '''
    row_item = kwargs.get("row_item")

    if not row_item:
        row_item = content

    text_list = get_product_text(is_content_item, content, row_item, col_headers)
    row_cells = table.add_row().cells
    element = content.element

    if (content.type == "Subtotal") or (content.type == "Total"):
        #Resets this var after use so the row header is correct
        row_item = content

    try:
        for (i, cell) in enumerate(row_cells):

            if not full_width_content(content, structure):
                add_content(cell, i, row_item, text_list, element, content=content)
                add_styling(cell, i, element)
            else:
                if i == 0:
                    add_content(cell, i, row_item, text_list, element)
                    add_styling(cell, i, element)
                else:
                    row_cells[0].merge(cell)

    except Exception as e:
        print(e)


def get_num_cols(structure):

    num_cols = 6

    if structure.hide_percentage:
        num_cols -= 1

    if structure.hide_ICR:
        num_cols -= 1

    return num_cols


def get_column_headers(structure):
    '''
    Gets the column headers for the product table.
    '''
    cols = ["Current", "Change", "Recommended"]

    if not structure.hide_percentage:
        cols.append("Allocation")

    if not structure.hide_ICR:
        cols.append("ICR")

    return cols

def get_product_text(is_content_item, content, row_item, col_headers):
    '''
    Returns a text_list to be inserted into each cell in the row.
    '''
    text_list = []

    if is_content_item:
        text_list = get_investment_data(content, row_item, col_headers)
    else:
        if content.type == 'Description':
            text_list = get_description_list(content, col_headers)
        elif "Subtotal" in content.type:
            text_list = get_product_subtotal_list(content, row_item)
        elif content.type == 'Total':
            text_list = get_product_total_list(content, row_item)
    return text_list

def get_investment_data(content, inv, col_headers):
    '''
    Currently gets all required data, which is then rejected if they don't fit
    the cols. May be changed in future to be more efficient.
    '''
    text_list = []

    text_list.append(inv.cur_value)
    text_list.append(inv.transaction)
    text_list.append(inv.rec_value)
    text_list.append(convert_to_percentage(get_investment_allocation(inv)))
    text_list.append(convert_to_percentage(inv.investment_fee))

    return text_list

def get_investment_allocation(inv):
    platform = inv.platform

    if platform.get_destination_platforms():
        return Decimal(0)
    else:
        return inv.allocation

def convert_to_percentage(item):
    '''
    Converts a percentage decimal to a readable percentage with % sign.
    '''
    return str(round((item * 100), 2))+'%'

def get_description_list(content, col_headers):
    '''
    Gets the column headers for the table.

    Currently just outputs the column headers, may change if customisable
    column headers allowed in future.
    '''
    return col_headers

def get_product_subtotal_list(content, row_item):
    platform = row_item
    statuses = ["cur_value", "transaction", "rec_value"]
    result = []
    for status in statuses:
        val = getattr(platform,"platform_"+status)
        result.append(val)
    result.extend(["",""])
    return result


def get_product_total_list(content, row_item):
    platforms = row_item
    statuses = ["cur_value", "transaction", "rec_value"]
    result = []
    for status in statuses:
        val = sum(getattr(p,"platform_"+status) for p in platforms)
        result.append(val)
    result.extend(["",""])
    return result


def remove_alternatives(all_platforms):
    return [p for p in all_platforms if p.status != "Alternative"]
