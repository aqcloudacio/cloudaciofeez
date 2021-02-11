import matplotlib as mpl
mpl.use('Agg')
import matplotlib.ticker as mtick
import pandas as pd
import seaborn as sns
sns.set()

from io import BytesIO


from django.shortcuts import get_list_or_404
from docx.enum.text import WD_BREAK

from documents.builder.styling import (add_styling,
                                       format_table)
from documents.builder.content import add_content
from documents.builder.structure import full_width_content
from documents.builder.utils import is_content_item

from investments.models import AssetAllocation


def get_aa(platforms, scenario):
    '''
    Collects the asset allocation of both current and recommended platforms.

    Then adds this to the RiskProfileAANames list for access in the doc builder.
    '''
    aa_list = get_list_or_404(AssetAllocation, investment__platform__in=platforms)
    cur = [aa for aa in aa_list if aa.investment.platform.status == "Current"]
    rec = [aa for aa in aa_list if aa.investment.platform.status == "Recommended"]

    curtotal = sum(p.platform_total for p in platforms if p.status == "Current")
    rectotal = sum(p.platform_total for p in platforms if p.status == "Recommended")

    rp = scenario.risk_profile
    rp_aa_list = rp.allocations.all()
    rp_aa_names = rp.group.rp_aa_names.all()

    for rp_name in rp_aa_names:
        rp_name.cur = 0
        rp_name.rec = 0
        rp_name.rp = 0

        for aa in cur:
            links = len(aa.name.rp_aaname_link.all())
            if rp_name in aa.name.rp_aaname_link.all():
                rp_name.cur += (aa.platform_percentage * (aa.investment.platform.platform_total / curtotal)) / links

        for aa in rec:
            links = len(aa.name.rp_aaname_link.all())
            if rp_name in aa.name.rp_aaname_link.all():
                rp_name.rec += (aa.platform_percentage * (aa.investment.platform.platform_total / rectotal)) / links

        for aa in rp_aa_list:
            if aa.name == rp_name:
                rp_name.rp += aa.percentage


    rp_aa_names = consolidate_names(rp_aa_names)

    return rp_aa_names


def consolidate_names(names):
    '''
    Consolidates duplicate custom_names in the rp_aa_names set.

    I DONT THINK THIS WORKS
    '''
    names = list(names)
    for name in names:
        unique_list = [x for x in names if ((x.custom_name == name.custom_name) and (x.id != name.id))]
        if unique_list:
            for item in unique_list:

                name.cur += item.cur
                name.rec += item.rec
                name.rp += item.rp

                names.pop(names.index(item))

    return names


def get_aa_table_types(structure):
    '''
    Returns the four types of AA tables that the user can include if desired.
    '''
    lst = []
    if structure.include_cur_vs_rp:
        lst.append("cur")
    if structure.include_rec_vs_rp:
        lst.append("rec")
    if structure.include_cur_vs_rp_vs_rec:
        lst.append("overall")
    if structure.include_cur_vs_rec:
        lst.append("cur_vs_rec")

    return lst

def get_num_cols(table_type, structure):

    if table_type == "overall":
        num_cols = 5
    else:
        num_cols = 4

    if structure.hide_variance:
        num_cols -= 1

    return num_cols


def build_aa_tables(document, style, table_style, structure, content, asset_allocation):

    aa_table_types = get_aa_table_types(structure)

    for table_type in aa_table_types:

        col_headers = get_column_headers(table_type, structure)

        num_cols = get_num_cols(table_type, structure)
        table = document.add_table(rows=0, cols=num_cols)

        format_table(table, table_style, style)
        build_rows(table, structure, content, asset_allocation, table_type, col_headers)
        if not structure.hide_risk_profile:
            add_chart(document, asset_allocation, col_headers)

        p = document.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)


def build_rows(table, structure, content_list, asset_allocation, table_type, col_headers):
    '''
    Calls functions to create rows/cells.
    '''
    asset_allocation = get_variance(asset_allocation, col_headers)

    try:
        for content in content_list:
            if not is_content_item(content):
                build_cells(table, content, structure, asset_allocation, is_content_item(content), col_headers)
            else:
                aa_list = [aa.custom_name for aa in asset_allocation if aa.asset_type == content.type.split()[0]]
                for aa_name in aa_list:
                    build_cells(table, content, structure, asset_allocation, is_content_item(content), col_headers, aa_name=aa_name)
    except Exception as e:
        print(e)



def build_cells(table, content, structure, asset_allocation, is_content_item, col_headers, **kwargs):
    '''
    Creates rows and cells and calls functions to add text and styling
    '''
    aa_name = kwargs.get("aa_name")

    if not aa_name:
        aa_name = content

    text_list = get_aa_text(is_content_item, content, asset_allocation, aa_name, col_headers)
    row_cells = table.add_row().cells
    element = content.element

    try:
        for (i, cell) in enumerate(row_cells):

            if not full_width_content(content, structure):
                add_content(cell, i, aa_name, text_list, element)
                add_styling(cell, i, element)
            else:
                if i == 0:
                    add_content(cell, i, aa_name, text_list, element)
                    add_styling(cell, i, element)
                else:
                    row_cells[0].merge(cell)

    except Exception as e:
        print(e)


def add_chart(document, asset_allocation, col_headers):

    # Builds data set for table.
    aa_list = [aa.custom_name for aa in asset_allocation]
    data = []
    col_headers = [c for c in col_headers if c != "var"]
    for header in col_headers:
        grouping = []
        for aa in asset_allocation:
            grouping.append(getattr(aa, header)*100)

        data.append(grouping)

    # Used to insert the plot in the docx file without saving to disc.
    memfile = BytesIO()

    # Sets some styling
    mpl.rcParams['font.size'] = 9
    mpl.rcParams['font.sans-serif'] = "Calibri"
    mpl.rcParams['font.family'] = "sans-serif"
    mpl.rcParams['figure.dpi'] = 700

    # Sets background and colour palette in seaborn
    sns.set_context("paper")
    sns.set_palette("muted")

    # Makes the column headers user-friendly for the legend
    col_headers = get_status_list(col_headers)

    # Convert data to pandas DataFrame.
    df = pd.DataFrame(data, index=col_headers, columns=aa_list).T
    df = df.astype(float)

    plot = df.plot(kind="bar", grid=True, rot=0, figsize=(7, 3.5))

    # Formats gridlines
    plot.yaxis.set_major_formatter(mtick.PercentFormatter())
    plot.grid(axis="x")
    plot.set_axisbelow(True)
    plot.spines['top'].set_visible(False)
    plot.spines['right'].set_visible(False)
    plot.spines['left'].set_visible(False)
    plot.axis('tight')

    fig = plot.get_figure()

    # bbox_inches hides whitespace
    fig.savefig(memfile, bbox_inches='tight')

    document.add_picture(memfile)
    # document.add_picture(memfile, width=Inches(1.25))
    memfile.close()


def get_aa_text(is_content_item, content, asset_allocation, aa_name, col_headers):
    '''
    Returns a text_list to be inserted into each cell in the row.
    '''
    text_list = []

    if is_content_item:
        text_list = get_aa_list(content, asset_allocation, aa_name, col_headers)
    else:
        if content.type == 'Status':
            text_list = get_status_list(col_headers)
        elif "Subtotal" in content.type:
            text_list = get_aa_subtotal_list(content, asset_allocation, col_headers)
            text_list = convert_to_percentage(text_list)
        elif content.type == 'Total':
            text_list = ["100.00%","100.00%","100.00%","0.00%"]
    return text_list

def get_status_list(col_headers):
    '''
    Gets the status/column headers for the table.
    '''
    header_names = []
    for header in col_headers:
        if header == "cur":
            header_names.append("Current")
        elif header == "rec":
            header_names.append("Recommended")
        elif header == "rp":
            header_names.append("Risk Profile")
        elif header == "var":
            header_names.append("Variance")

    return header_names

def get_aa_subtotal_list(content, asset_allocation, col_headers):
    '''
    Gets a subtotal list based on the given column headers and specific content.
    '''
    status = content.type.split()[0]
    dict = {}
    for header in col_headers:
        dict[header] = sum(getattr(aa,header) for aa in asset_allocation if aa.asset_type == status)

    return list(dict.values())


def get_variance(asset_allocation, col_headers):
    '''
    Calculates the variance for each asset allocation item. Varies and is
    recalculated for each aa table type.
    '''
    for aa in asset_allocation:
        aa.var = 0
        if "rp" in col_headers:
            if "rec" not in col_headers:
                aa.var = aa.cur - aa.rp
            else:
                aa.var = aa.rec - aa.rp
        else:
            aa.var = aa.rec - aa.cur

    return asset_allocation

def get_aa_list(content, asset_allocation, aa_name, col_headers):
    '''
    Searches asset allocation objects for a attribute that matches the header.
    '''
    text_list = []
    status = content.type.split()[0]

    try:
        for aa in asset_allocation:
            if (aa.custom_name == aa_name) and (aa.asset_type == status):
                for header in col_headers:
                    text_list.append(getattr(aa, header))

                text_list = convert_to_percentage(text_list)
                return text_list
                break
    except Exception as e:
        print(e)


def convert_to_percentage(list):
    '''
    Converts a percentage decimal to a readable percentage with % sign.
    '''
    for index, item in enumerate(list):
        list[index] = str(round((item * 100), 2))+'%'
    return list


def get_column_headers(table_type, structure):
    '''
    Gets the column headers for each type of table. Removes "Var" if that
    option has been selected in the structure model.
    '''
    cols = []

    if table_type == "cur":
        cols = ["cur", "rp", "var"]
    elif table_type == "rec":
        cols = ["rec", "rp", "var"]
    elif table_type == "overall":
        cols = ["cur", "rp", "rec", "var"]
    elif table_type == "cur_vs_rec":
        cols = ["cur", "rec", "var"]

    if structure.hide_variance:
        cols.remove("var")

    return cols
