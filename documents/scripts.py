from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table



def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    # accesses the actual element
    tcPr = tc.get_or_add_tcPr()
    # print(dir(tc))
    # i think access the table paragraph element
    tcMar = OxmlElement('w:tcMar')
    # access the table cell margin element

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            # w:top
            node.set(qn('w:w'), str(kwargs.get(m)))
            # w:w set kwargs[0]
            node.set(qn('w:type'), 'dxa')
            # w:type set 'dxa'
            tcMar.append(node)

    tcPr.append(tcMar)

############################################


def set_table_cell_margins(table: Table, **kwargs):
    """
    table:  actual table instance you want to modify

    usage:

        set_table_cell_margins(table, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tbl = table._tbl
    print(dir(table))
    # accesses the actual element
    tblPr = table._tblPr
    # i think access the table paragraph element
    tblCellMar = OxmlElement('w:tblCellMar')
    # access the table cell margin element

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            # w:top
            node.set(qn('w:w'), str(kwargs.get(m)))
            # w:w set kwargs[0]
            node.set(qn('w:type'), 'dxa')
            # w:type set 'dxa'
            tblCellMar.append(node)

    tblPr.append(tblCellMar)

########################################

document = Document()

document.add_heading('Document Title', 0)


# records = (
#     ("- Administration Fee", '$101.01', '$0.00'),
#     ("- Expense Recovery", '$3.33', '$0.00'),
#     ("- Trustee Fee", '$555.45', '$444.44')
# )

records = (
    ("- Administration Fee"),
    ("- Expense Recovery"),
    ("- Trustee Fee")
)

platforms = ['BT Panorama Super', "BT Wrap"]


table = document.add_table(rows=1, cols=1)

set_table_cell_margins(table, top=56.693, start=113.389, bottom=56.693, end=113.389)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Description'
# hdr_cells[1].text = 'Current'
# hdr_cells[2].text = 'Proposed'
for description in records:
    row_cells = table.add_row().cells
    celly = row_cells[0]
    celly.text = description
    celly.paragraphs[0].runs[0].font.size = Pt(10)
    celly.paragraphs[0].runs[0].font.name = 'Calibri'
    set_cell_margins(celly, top=56.693, start=113.389, bottom=56.693, end=113.389)

    #
    # row_cells[1].text = fee1
    # row_cells[2].text = fee2

# test_item = row_cells[0].paragraphs[0].runs[0].font
#
# test_item
# print(test_item)

# document.save('demo.docx')
